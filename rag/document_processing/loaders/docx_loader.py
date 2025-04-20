"""
DOCX document loader for the RAG system.
"""
import os
from typing import Dict, List, Optional, Any
from pathlib import Path
import docx
from langchain.document_loaders import Docx2txtLoader
from langchain.schema import Document

class DOCXLoader:
    """Loader for DOCX documents."""
    
    def __init__(self, use_langchain: bool = True):
        """
        Initialize the DOCX loader.
        
        Args:
            use_langchain: Whether to use LangChain's loader or a custom implementation
        """
        self.use_langchain = use_langchain
    
    def load(self, file_path: str) -> List[Document]:
        """
        Load a DOCX document.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            List of Document objects
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        if self.use_langchain:
            return self._load_with_langchain(file_path)
        else:
            return self._load_custom(file_path)
    
    def _load_with_langchain(self, file_path: str) -> List[Document]:
        """Load DOCX using LangChain's loader."""
        loader = Docx2txtLoader(file_path)
        documents = loader.load()
        
        # Add additional metadata
        for doc in documents:
            doc.metadata["file_path"] = file_path
            doc.metadata["file_type"] = "docx"
            doc.metadata["file_name"] = os.path.basename(file_path)
        
        return documents
    
    def _load_custom(self, file_path: str) -> List[Document]:
        """Custom DOCX loading implementation."""
        documents = []
        
        try:
            doc = docx.Document(file_path)
            
            # Extract metadata
            metadata = {
                "file_path": file_path,
                "file_type": "docx",
                "file_name": os.path.basename(file_path),
                "total_paragraphs": len(doc.paragraphs)
            }
            
            # Extract core properties if available
            try:
                core_props = doc.core_properties
                if core_props:
                    for prop in ['author', 'category', 'comments', 'content_status', 
                                'created', 'identifier', 'keywords', 'language', 
                                'last_modified_by', 'last_printed', 'modified', 
                                'revision', 'subject', 'title', 'version']:
                        if hasattr(core_props, prop):
                            value = getattr(core_props, prop)
                            if value:
                                metadata[f"docx_{prop}"] = str(value)
            except:
                pass
            
            # Extract text from paragraphs
            full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            if full_text.strip():
                documents.append(Document(
                    page_content=full_text,
                    metadata=metadata
                ))
        
        except Exception as e:
            raise Exception(f"Error loading DOCX {file_path}: {e}")
        
        return documents